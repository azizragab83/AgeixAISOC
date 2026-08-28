# -*- coding: utf-8 -*-
"""Add the official Title Page to the existing final thesis docx."""
import os
import shutil
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

DOC_DIR = r"C:\Users\Digilians\Downloads\doc"
SRC = os.path.join(DOC_DIR, "AegisX_SOC_Thesis_Final.docx")
DST = SRC  # in-place update after backup
BAK = os.path.join(DOC_DIR, "_backup_thesis_no_titlepage.docx")

TITLE = "AgeixAISOC: An AI-Orchestrated Cyber Defense Platform with Human-in-the-Loop Governance"
TEAM = ["Aziz Ragab Aziz (Team Leader)", "Mohamed Hany", "Emad Hassan", "Taha Elghonaimi"]
SUP = "Under the supervision of"
SUP_NAME = "Dr. Rabab M. Nabawy  -  General Supervisor"
SUP2 = "Dr. Ahmed Tobal  -  Academic Director"
DATE = "August 2026"


def add_center(doc, text, size, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    f = r.font
    f.name = "Times New Roman"
    f.size = Pt(size)
    f.bold = bold
    f.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def main():
    shutil.copyfile(SRC, BAK)
    d = docx.Document(SRC)

    from docx.text.paragraph import Paragraph as P
    first_p_elem = d.paragraphs[0]._p
    anchor_par = P(first_p_elem, d)  # Paragraph wrapper for first paragraph

    def insert_before(runs=None, align=None):
        """Insert new paragraph(s) before the first paragraph."""
        new_p = anchor_par.insert_paragraph_before("")
        if runs:
            for txt, size, bold, ital in runs:
                r = new_p.add_run(txt)
                rf = r.font
                rf.name = "Times New Roman"
                rf.size = Pt(size)
                rf.bold = bold
                rf.italic = ital
        if align is not None:
            new_p.alignment = align
        return new_p

    C = WD_ALIGN_PARAGRAPH.CENTER

    lines = [
        ("MSA University", 16, True, False),
        ("Digilians 9-Month Diploma in Cybersecurity", 14, False, True),
        ("Cybersecurity Track", 14, False, True),
        ("", 12, False, False),
        ("Graduation Project Document", 16, True, False),
        ("", 12, False, False),
    ]
    for txt, size, bold, ital in lines:
        insert_before(runs=[(txt, size, bold, ital)], align=C)

    # Project title
    insert_before(runs=[(TITLE, 20, True, False)], align=C)

    more = [
        ("", 12, False, False),
        ("A thesis submitted by", 14, False, True),
    ]
    for txt, size, bold, ital in more:
        insert_before(runs=[(txt, size, bold, ital)], align=C)

    for name in TEAM:
        insert_before(runs=[(name, 14, True, False)], align=C)

    sup_lines = [
        ("", 12, False, False),
        (SUP, 14, True, True),
        (SUP_NAME, 14, True, False),
        (SUP2, 14, True, False),
        ("", 12, False, False),
        (DATE, 14, True, False),
    ]
    for txt, size, bold, ital in sup_lines:
        insert_before(runs=[(txt, size, bold, ital)], align=C)

    # Page break so Approval Sheet starts on next page
    brk = anchor_par.insert_paragraph_before("")
    run = brk.add_run()
    run.add_break(WD_BREAK.PAGE)

    d.save(DST)
    print("[OK] title page added ->", DST)


if __name__ == "__main__":
    main()
