"""Shared document setup and helpers for the AgeixAISOC thesis builder."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"C:\Users\Digilians\Downloads\docm"
OUT_PATH = os.path.join(OUT_DIR, "AgeixAISOC_Final_Thesis.docx")

doc = None


def init_document():
    global doc
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return doc


def add_para(text, bold=False, italic=False, align=None, size=12, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
    return h


def add_code_block(code):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    return p


def add_table(headers, rows, caption=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
    doc.add_paragraph()
    return t


def add_field(paragraph_text, instr):
    """Insert a Word field (TOC etc.) into the document."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = paragraph_text
    r2 = OxmlElement("w:r"); r2.append(t)
    fldChar3 = OxmlElement("w:fldChar"); fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar); run._r.append(instrText); run._r.append(fldChar2)
    run._r.append(r2); run._r.append(fldChar3)
    return p


def add_page_break():
    doc.add_page_break()